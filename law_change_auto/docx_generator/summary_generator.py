"""법령제·개정 주요내용 요약 문서(docx) 생성.

기간(또는 단일 일자) 동안 수집된 LawChangeDetail 목록을 받아
법무팀 배포용 "주요내용 요약" 1건을 생성한다. 구조:

  Ⅰ. 개요          — 유형별 건수
  Ⅱ. 당사 관련성 높은 주요 개정 사항 — 생명보험사 관련 키워드 매칭 건 표
  Ⅲ. 전체 법령 제·개정 요약          — 전체 건 표
  Ⅳ. 미결사항      — 입법예고/규정변경예고·미래 시행일 등 자동 추출

각 행의 "주요 내용 요약"은 LLM(Groq→Gemini) 한 줄 요약을 우선 사용하고,
키 부재·실패 시 본문(주요내용/개정이유) 추출 기반 기본 요약으로 대체한다.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

from ..models import LawChangeDetail
from .generator import guide_display_title

_FONT = "KoPub돋움체_Pro Light"
_FONT_BOLD = "KoPub돋움체_Pro Bold"

# 생명보험회사 업무 관련성이 높은 법령 키워드 (Ⅱ 표 선별용)
_RELEVANCE_KEYWORDS = (
    "보험", "신용정보", "자본시장", "금융투자", "금융복합", "금융지주",
    "개인정보", "전자금융", "여신전문", "근로자퇴직급여", "퇴직연금",
    "국민연금", "고용보험", "금융소비자", "특정금융거래", "자금세탁",
    "개인금융채권", "금융기관검사", "외부감사", "조세특례제한법", "공정거래",
)


def _set_run_font(run, *, bold: bool = False, size: int = 11) -> None:
    run.font.size = Pt(size)
    name = _FONT_BOLD if bold else _FONT
    run.font.name = name
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)


def _add_para(doc, text: str = "", *, bold: bool = False, size: int = 11,
              align=None, space_after: int = 6):
    p = doc.add_paragraph()
    if text:
        _set_run_font(p.add_run(text), bold=bold, size=size)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def _is_relevant(detail: LawChangeDetail) -> bool:
    name = detail.meta.law_name or ""
    return any(kw in name for kw in _RELEVANCE_KEYWORDS)


def _change_date_str(detail: LawChangeDetail) -> str:
    """표의 '시행/예고일' 칸 문자열."""
    meta = detail.meta
    if meta.category == "입법예고":
        label = "예고"
        d = meta.announcement_date
        deadline = detail.opinion_deadline
        if deadline:
            return f"{label} {deadline}"
    else:
        label = "시행"
        d = meta.effective_date or meta.announcement_date
    if not d:
        return "-"
    return f"{label} {d.year}.{d.month}.{d.day}."


def _short_name(detail: LawChangeDetail) -> str:
    """표 '법령명' 칸: 입법예고/규정변경예고는 구분 병기."""
    name = (detail.meta.law_name or "").strip()
    ct = detail.meta.change_type
    if detail.meta.category == "입법예고" and ct in ("입법예고", "규정변경예고"):
        # 이미 제목에 구분이 들어있지 않으면 병기
        if ct not in name:
            return f"{name} ({ct})"
    return name


def _strip_lead(text: str) -> str:
    text = " ".join((text or "").split())
    text = text.lstrip("1234567890.가나다라마바사아·◎- \t")
    # 문장 중간 머리표(가. 나. 다. / 1. 2.)는 쉼표로 치환해 한 줄로 정리
    text = re.sub(r"\s+(?:[가나다라마바사아]|\d{1,2})\.\s+", ", ", text)
    return text.strip(" ,")


def _fallback_summary(detail: LawChangeDetail) -> str:
    """LLM 미사용/실패 시 본문에서 한 줄 요약 추출."""
    pools: List[str] = []
    pools.extend(detail.main_change_sections or [])
    pools.extend(detail.combined_reason_and_main_sections or [])
    pools.extend(detail.reason_sections or [])
    for raw in pools:
        s = _strip_lead(raw)
        if len(s) >= 10:
            return (s[:120].rstrip() + "…") if len(s) > 120 else s
    # 신구조문만 있는 경우
    if detail.article_comparisons:
        return f"신·구조문 대비표 {len(detail.article_comparisons)}개 항목 개정"
    return "세부 개정내용은 안내서 본문 참조"


def _one_line_summary(detail: LawChangeDetail, *, use_llm: bool) -> str:
    if use_llm:
        try:
            from ..services.gemini_client import fetch_one_line_summary
            llm = fetch_one_line_summary(
                detail.meta.law_name or "",
                detail.reason_sections or detail.combined_reason_and_main_sections,
                detail.main_change_sections or detail.combined_reason_and_main_sections,
            )
            if llm:
                return llm
        except Exception:
            pass
    return _fallback_summary(detail)


def _add_summary_table(doc, items: List[Tuple[int, LawChangeDetail, str]]) -> None:
    """[(번호, detail, 요약문), ...]로 4열 표 작성."""
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("No.", "법령명", "시행/예고일", "주요 내용 요약")
    for cell, htext in zip(table.rows[0].cells, headers):
        cell.text = ""
        _set_run_font(cell.paragraphs[0].add_run(htext), bold=True, size=10)
    widths = (None,)
    for no, detail, summary in items:
        row = table.add_row().cells
        for cell, val, sz in (
            (row[0], str(no), 10),
            (row[1], _short_name(detail), 10),
            (row[2], _change_date_str(detail), 10),
            (row[3], summary, 10),
        ):
            cell.text = ""
            _set_run_font(cell.paragraphs[0].add_run(val), size=sz)


def _overview_line(details: List[LawChangeDetail]) -> str:
    law = sum(1 for d in details if d.meta.category == "법령")
    admin = sum(1 for d in details if d.meta.category == "행정규칙")
    legis = sum(1 for d in details if d.meta.category == "입법예고"
                and d.meta.change_type != "규정변경예고")
    reg = sum(1 for d in details if d.meta.change_type == "규정변경예고")
    parts = [f"법령 {law}건"]
    if admin:
        parts.append(f"행정규칙 {admin}건")
    if reg:
        parts.append(f"규정변경예고 {reg}건")
    if legis:
        parts.append(f"입법예고 {legis}건")
    return ", ".join(parts)


def _brief_name(detail: LawChangeDetail, limit: int = 22) -> str:
    """미결사항 나열용 짧은 법령명. 제목에 든 괄호를 제거하고 깔끔히 절단."""
    name = re.sub(r"[「」｢｣（）()]", "", detail.meta.law_name or "").strip()
    name = re.sub(r"\s+", " ", name)
    if len(name) > limit:
        name = name[:limit].rstrip() + "…"
    return f"「{name}」"


def _open_issues(details: List[LawChangeDetail], guide_date: date) -> List[str]:
    """미결사항 자동 추출."""
    issues: List[str] = []
    pending = [d for d in details if d.meta.category == "입법예고"
               or d.meta.change_type == "규정변경예고"]
    if pending:
        names = ", ".join(_brief_name(d) for d in pending[:4])
        more = f" 외 {len(pending) - 4}건" if len(pending) > 4 else ""
        issues.append(
            f"입법예고·규정변경예고 {len(pending)}건({names}{more})은 확정 시 내용이 "
            f"변경될 수 있으므로, 최종 확정 후 재검토가 필요함."
        )
    future = [d for d in details
              if d.meta.effective_date and d.meta.effective_date > guide_date]
    if future:
        names = ", ".join(_brief_name(d) for d in future[:4])
        more = f" 외 {len(future) - 4}건" if len(future) > 4 else ""
        issues.append(
            f"시행일 미도래 {len(future)}건({names}{more})은 시행 전 하위법령·"
            f"세부지침 동향을 지속 모니터링할 필요가 있음."
        )
    relevant = [d for d in details if _is_relevant(d)]
    if relevant:
        issues.append(
            "당사 관련성 높은 개정사항은 소관 부서 검토를 거쳐 내규·약관·실무 "
            "절차에 반영할 필요가 있음."
        )
    issues.append("각 건의 세부 개정내용 및 신·구조문 대비표는 개별 안내서 본문을 참조 바람.")
    return issues


def generate_summary_docx(
    details: List[LawChangeDetail],
    output_path: Path,
    *,
    period_line: str,
    guide_date: date,
    use_llm: bool = True,
) -> Optional[Path]:
    """주요내용 요약 docx 1건 생성. 생성 실패·대상 없음 시 None."""
    details = [d for d in details if d.has_any_content() or d.article_comparisons]
    if not details:
        return None

    # 본문 요약문 사전 계산 (LLM 호출 1회/건)
    summaries = [_one_line_summary(d, use_llm=use_llm) for d in details]

    doc = Document()
    # 표지
    _add_para(doc, f"{period_line} 법령제·개정 주요내용 요약",
              bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_para(doc, f"{guide_date.year}. {guide_date.month:02d}.",
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_para(doc, "법 무 팀", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Ⅰ. 개요
    _add_para(doc, "Ⅰ. 개요", bold=True, size=12, space_after=4)
    _add_para(
        doc,
        f"{period_line} 기간 중 모니터링 대상 법령에서 제·개정·예고된 사항을 "
        f"유형별로 분류·요약함. 총 {len(details)}건"
        f"({_overview_line(details)}).",
        space_after=10,
    )

    # Ⅱ. 당사 관련성 높은 주요 개정 사항
    relevant_items = [
        (i + 1, d, summaries[i]) for i, d in enumerate(details) if _is_relevant(d)
    ]
    _add_para(doc, "Ⅱ. 당사 관련성 높은 주요 개정 사항", bold=True, size=12, space_after=4)
    if relevant_items:
        _add_summary_table(doc, relevant_items)
    else:
        _add_para(doc, "해당 기간 중 당사 관련성이 두드러지는 개정 사항은 없음.")
    _add_para(doc, "", space_after=8)

    # Ⅲ. 전체 법령 제·개정 요약
    all_items = [(i + 1, d, summaries[i]) for i, d in enumerate(details)]
    _add_para(doc, f"Ⅲ. 전체 법령 제·개정 요약 ({len(details)}건)",
              bold=True, size=12, space_after=4)
    _add_summary_table(doc, all_items)
    _add_para(doc, "", space_after=8)

    # Ⅳ. 미결사항
    _add_para(doc, "Ⅳ. 미결사항", bold=True, size=12, space_after=4)
    for issue in _open_issues(details, guide_date):
        p = doc.add_paragraph(style="List Bullet")
        _set_run_font(p.add_run(issue), size=11)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
