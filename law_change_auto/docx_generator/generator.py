from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import List, Tuple

# 개정이유 본문에서 목차와 중복되는 접두어 제거 (예: [일부개정], ◇ 개정이유 및 주요내용)
_REVISION_PREFIXES = re.compile(
    r"^\s*(\[일부개정\]|\[전부개정\]|\[타법개정\]|\[제정\]|◇\s*개정이유\s*및\s*주요내용|◇\s*개정이유|◇\s*주요내용)\s*\n?"
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from ..models import LawChangeDetail, LawChangeDetailSeq


LINE_SPACING_BODY = 1.5
LINE_SPACING_TITLE = 1.0
PARAGRAPH_SPACING_BEFORE = Pt(0)
PARAGRAPH_SPACING_AFTER = Pt(6)


def _apply_body_format(paragraph) -> None:
    paragraph.paragraph_format.line_spacing = LINE_SPACING_BODY
    paragraph.paragraph_format.space_before = PARAGRAPH_SPACING_BEFORE
    paragraph.paragraph_format.space_after = PARAGRAPH_SPACING_AFTER


class DocxGenerator:
    """legal_doc_converter의 DocxGenerator 형식을 따른 문서 생성기."""

    def __init__(self) -> None:
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        margin = Mm(12.7)
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    def _setup_styles(self) -> None:
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "KoPub돋움체_Pro Light"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "KoPub돋움체_Pro Light")
        style.paragraph_format.line_spacing = LINE_SPACING_BODY

    def add_title(self, title_text: str) -> None:
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "KoPub돋움체_Pro Bold"
        title.paragraph_format.line_spacing = LINE_SPACING_TITLE
        title.paragraph_format.space_after = Pt(0)

    def add_metadata(
        self,
        enforcement_date: str,
        law_number: str,
        amendment_date: str,
        amendment_type: str = "",
        law_type_label: str = "법률",
        date_str: str = "25. 01.",
        dept: str = "법 무 팀",
    ) -> None:
        if enforcement_date or law_number or amendment_date:
            meta1 = self.doc.add_paragraph()
            meta1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if amendment_type:
                meta_line = f"[시행 {enforcement_date}] [{law_type_label} 제{law_number}호, {amendment_date}, {amendment_type}]"
            else:
                meta_line = f"[시행 {enforcement_date}] [{law_type_label} 제{law_number}호, {amendment_date}]"
            run = meta1.add_run(meta_line)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "KoPub돋움체_Pro Bold"
            meta1.paragraph_format.line_spacing = LINE_SPACING_TITLE
            meta1.paragraph_format.space_before = PARAGRAPH_SPACING_BEFORE
            meta1.paragraph_format.space_after = PARAGRAPH_SPACING_AFTER

        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.add_run(date_str)
        _apply_body_format(date_p)

        dept_p = self.doc.add_paragraph()
        dept_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dept_p.add_run(dept)
        _apply_body_format(dept_p)

        self.doc.add_paragraph()

    def add_notice_metadata(
        self,
        notice_period: str,
        law_type_name: str = "",
        date_str: str = "25. 01.",
        dept: str = "법 무 팀",
    ) -> None:
        """입법예고용 메타데이터 (예고기간)."""
        if notice_period:
            meta1 = self.doc.add_paragraph()
            meta1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_line = f"[예고기간: {notice_period}]"
            run = meta1.add_run(meta_line)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "KoPub돋움체_Pro Bold"
            meta1.paragraph_format.line_spacing = LINE_SPACING_TITLE
            meta1.paragraph_format.space_before = PARAGRAPH_SPACING_BEFORE
            meta1.paragraph_format.space_after = PARAGRAPH_SPACING_AFTER

        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.add_run(date_str)
        _apply_body_format(date_p)

        dept_p = self.doc.add_paragraph()
        dept_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dept_p.add_run(dept)
        _apply_body_format(dept_p)

        self.doc.add_paragraph()

    def add_section(
        self, number: str, title: str, content: str | List[str] | None = None, is_bold: bool = False
    ) -> None:
        section = self.doc.add_paragraph()
        run = section.add_run(f"{number}. {title}")
        run.bold = True
        run.font.size = Pt(11)
        _apply_body_format(section)

        if content:
            if isinstance(content, list):
                for para_text in content:
                    text = para_text.strip()
                    if not text:
                        continue
                    content_p = self.doc.add_paragraph()
                    content_run = content_p.add_run(text)
                    if is_bold:
                        content_run.bold = True
                    content_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    content_p.paragraph_format.first_line_indent = Inches(0)
                    _apply_body_format(content_p)
            else:
                text = content.strip()
                if text:
                    content_p = self.doc.add_paragraph()
                    content_run = content_p.add_run(text)
                    if is_bold:
                        content_run.bold = True
                    content_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    _apply_body_format(content_p)

        self.doc.add_paragraph()

    def add_main_contents(self, intro_paragraphs: List[str] | None = None) -> None:
        if intro_paragraphs:
            for para_text in intro_paragraphs:
                text = para_text.strip()
                if not text:
                    continue
                p = self.doc.add_paragraph()
                p.add_run(text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _apply_body_format(p)
            self.doc.add_paragraph()

        self.doc.add_paragraph()

    def add_comparison_table(self, comparison_data: List[Tuple[str, str]]) -> None:
        if not comparison_data:
            return

        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "개정 전"
        header_cells[1].text = "개정 후"

        for cell in header_cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_background(cell, "D9D9D9")

        for old_text, new_text in comparison_data:
            row_cells = table.add_row().cells
            row_cells[0].text = (old_text or "").strip()
            row_cells[1].text = (new_text or "").strip()

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = LINE_SPACING_BODY
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        self.doc.add_paragraph()

    def _set_cell_background(self, cell, color: str) -> None:
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def save(self, filename: str) -> None:
        self.doc.save(filename)


def _clean_revision_paras(paras: List[str]) -> List[str]:
    """개정이유 본문에서 불필요한 접두어를 제거하고, 개행 기준으로 문단을 나눈다."""
    result: List[str] = []
    for block in paras:
        text = block.strip()
        while True:
            m = _REVISION_PREFIXES.match(text)
            if not m:
                break
            text = text[m.end() :].strip()
        for line in text.split("\n"):
            line = line.strip()
            if line:
                result.append(line)
    return result


def _detail_to_docx(detail: LawChangeDetail, target_date: date, generator: DocxGenerator) -> None:
    meta = detail.meta

    if meta.category == "행정규칙":
        title_suffix = "고시 규정변경예고 안내"
    elif meta.category == "입법예고":
        title_suffix = "입법예고 안내"
    else:
        title_suffix = "시행 안내"

    generator.add_title(f"{meta.law_name} {title_suffix}")

    def _fmt_date(d: date | None) -> str:
        if not d:
            return ""
        return f"{d.year}. {d.month}. {d.day}."

    if meta.category == "입법예고":
        # 입법예고: 예고기간 표시
        start_str = _fmt_date(meta.announcement_date)
        end_str = _fmt_date(meta.effective_date)
        notice_period = f"{start_str} ~ {end_str}" if start_str and end_str else ""
        generator.add_notice_metadata(
            notice_period=notice_period,
            date_str=target_date.strftime("%Y. %m."),
        )
    else:
        enforcement_date = _fmt_date(meta.effective_date)
        amendment_date = meta.amendment_date_str or _fmt_date(meta.announcement_date)
        law_number = meta.law_number or meta.law_id or ""
        amendment_type = meta.amendment_type or ""

        # 법령 유형 레이블 결정: API 메타 > 법령명 기반 추론 > 기본값
        law_type_label = meta.law_type_label
        if not law_type_label:
            name = meta.law_name
            if "시행규칙" in name:
                law_type_label = "부령"
            elif "시행령" in name:
                law_type_label = "대통령령"
            elif meta.category == "행정규칙":
                law_type_label = "고시"
            else:
                law_type_label = "법률"

        generator.add_metadata(
            enforcement_date=enforcement_date,
            law_number=law_number,
            amendment_date=amendment_date,
            amendment_type=amendment_type,
            law_type_label=law_type_label,
            date_str=target_date.strftime("%Y. %m."),
        )

    is_combined = bool(detail.combined_reason_and_main_sections)

    if is_combined:
        paras = _clean_revision_paras(detail.combined_reason_and_main_sections or [])
        if paras:
            generator.add_section("1", "개정이유 및 주요내용", paras)
        else:
            generator.add_section("1", "개정이유 및 주요내용", "")
    else:
        reason_paras = _clean_revision_paras(detail.reason_sections or [])
        if reason_paras:
            generator.add_section("1", "개정이유", reason_paras)
        else:
            generator.add_section("1", "개정이유", "")

        main_paras = _clean_revision_paras(detail.main_change_sections or [])
        generator.add_section("2", "주요내용", main_paras or None)

    if detail.impact_analysis:
        impact_text = detail.impact_analysis
    else:
        impact_text = f"{meta.law_name} 개정에 따른 실무 영향을 면밀히 검토하여 관련 업무에 반영 바람."

    is_notice = meta.category == "입법예고"

    if is_combined:
        sec_impact = "2"
        sec_table = "3"
    else:
        sec_impact = "3"
        sec_table = "4"

    generator.add_section(sec_impact, "파급효과", impact_text, is_bold=True)

    if is_notice:
        # 입법예고: 향후 일정
        next_num = str(int(sec_impact) + 1)
        end_str = _fmt_date(meta.effective_date) if meta.effective_date else ""
        schedule_text = f"의견제출 마감: {end_str}" if end_str else "의견제출 기간 확인 필요"
        generator.add_section(next_num, "향후 일정", schedule_text)

    else:
        generator.add_section(sec_table, "신구조문 대비표")
        comparison_table: List[Tuple[str, str]] = [
            ((row.old_text or "").strip(), (row.new_text or "").strip()) for row in detail.article_comparisons
        ]
        generator.add_comparison_table(comparison_table)


def generate_guide(details: LawChangeDetailSeq, target_date: date, output_path: Path) -> Path:
    """LawChangeDetail를 legal_doc_converter 스타일 DOCX로 변환한다.

    전달된 모든 detail을 문서에 포함한다.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    generator = DocxGenerator()

    if details:
        for i, detail in enumerate(details):
            if i > 0:
                # 법령 간 구분을 위한 페이지 구분
                generator.doc.add_page_break()
            _detail_to_docx(detail, target_date, generator)
    else:
        generator.add_title("법령제·개정 안내서")
        generator.add_metadata("", "", "", date_str=target_date.strftime("%y. %m."))

    generator.save(str(output_path))
    return output_path

